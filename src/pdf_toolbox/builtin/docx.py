"""Convert PDF files to DOCX documents."""

from __future__ import annotations

import io
from pathlib import Path
from threading import Event

import fitz  # type: ignore  # pdf-toolbox: PyMuPDF lacks type hints | issue:-
from docx import Document
from PIL import Image

from pdf_toolbox.actions import action
from pdf_toolbox.utils import (
    logger,
    open_pdf,
    raise_if_cancelled,
    sane_output_dir,
)

RGB_COMPONENTS = 3


@action(category="Office")
def pdf_to_docx(
    input_pdf: str,
    out_dir: str | None = None,
    cancel: Event | None = None,
) -> str:
    """Convert a PDF into a DOCX document.

    The text of each page is extracted and appended to a Word document. Images
    encountered in the PDF are embedded as PNGs in the corresponding position.
    The resulting file is stored next to ``input_pdf`` unless ``out_dir`` points
    to a different directory. The path to the created DOCX file is returned.
    """
    logger.info("Converting %s to DOCX", input_pdf)
    docx_doc = Document()
    with open_pdf(input_pdf) as pdf:
        for page in pdf:
            raise_if_cancelled(cancel)
            text = page.get_text()
            if text:
                docx_doc.add_paragraph(text)
            for img in page.get_images(full=True):
                raise_if_cancelled(cancel)
                xref = img[0]
                pix = fitz.Pixmap(pdf, xref)
                if (
                    pix.n > RGB_COMPONENTS
                ):  # pragma: no cover  # pdf-toolbox: rare colorspace | issue:-
                    pix = fitz.Pixmap(
                        fitz.csRGB, pix
                    )  # pragma: no cover  # pdf-toolbox: rare colorspace | issue:-
                pil = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                with io.BytesIO() as buf:
                    pil.save(buf, format="PNG")
                    buf.seek(0)
                    docx_doc.add_picture(buf)
    out_path = sane_output_dir(input_pdf, out_dir) / f"{Path(input_pdf).stem}.docx"
    docx_doc.save(str(out_path))
    logger.info("DOCX written to %s", out_path)
    return str(out_path)
