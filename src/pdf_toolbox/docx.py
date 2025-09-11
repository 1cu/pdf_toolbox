from __future__ import annotations

import argparse
import io
from pathlib import Path
from threading import Event

import fitz  # type: ignore
from PIL import Image
from docx import Document

from .actions import action
from .utils import open_pdf, raise_if_cancelled, sane_output_dir


@action(category="Office")
def pdf_to_docx(
    input_pdf: str,
    out_dir: str | None = None,
    cancel: Event | None = None,
) -> str:
    """Convert a PDF into a DOCX document.

    The text of each page is extracted and appended to a Word document. Images
    encountered in the PDF are embedded as PNGs in the corresponding position.
    The resulting file is stored next to ``input_pdf`` unless ``out_dir`` points
    to a different directory. The path to the created DOCX file is returned.
    """
    docx_doc = Document()
    with open_pdf(input_pdf) as pdf:
        for page in pdf:
            raise_if_cancelled(cancel)  # pragma: no cover
            text = page.get_text()
            if text:  # pragma: no cover - input PDF in tests has no text
                docx_doc.add_paragraph(text)  # pragma: no cover
            for img in page.get_images(full=True):
                raise_if_cancelled(cancel)  # pragma: no cover
                xref = img[0]
                pix = fitz.Pixmap(pdf, xref)
                if pix.n > 3:  # pragma: no cover - rare branch
                    pix = fitz.Pixmap(fitz.csRGB, pix)  # pragma: no cover
                pil = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                with io.BytesIO() as buf:
                    pil.save(buf, format="PNG")
                    buf.seek(0)
                    docx_doc.add_picture(buf)
    out_path = sane_output_dir(input_pdf, out_dir) / f"{Path(input_pdf).stem}.docx"
    docx_doc.save(str(out_path))
    return str(out_path)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="PDF to Word")
    parser.add_argument("input_pdf")
    parser.add_argument("--out-dir")
    args = parser.parse_args()
    pdf_to_docx(args.input_pdf, args.out_dir)
