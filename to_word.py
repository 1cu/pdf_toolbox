import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
import os
import sys
from tempfile import NamedTemporaryFile

def pdf_to_word(input_pdf):
    """Konvertiert eine PDF-Datei in ein Word-Dokument mit Text, Tabellen und Diagrammen als Bilder."""
    
    if not os.path.isfile(input_pdf):
        print(f"Fehler: Datei {input_pdf} nicht gefunden.")
        return

    base_dir, filename = os.path.split(input_pdf)
    base_name, _ = os.path.splitext(filename)
    output_docx = os.path.join(base_dir, f"{base_name}.docx")

    try:
        doc = fitz.open(input_pdf)
        word_doc = Document()

        for i, page in enumerate(doc, start=1):
            text_blocks = page.get_text("dict")["blocks"]

            for block in text_blocks:
                if "lines" in block:
                    for line in block["lines"]:
                        paragraph = word_doc.add_paragraph()
                        for span in line["spans"]:
                            run = paragraph.add_run(span["text"])
                            
                            # Formatierungen erkennen
                            if "bold" in span and span["bold"]:
                                run.bold = True
                            if "italic" in span and span["italic"]:
                                run.italic = True
                            if span["size"] > 12:  # Heuristik für Überschriften
                                paragraph.style = "Heading 1"

            # Tabellen extrahieren & ins Word-Dokument einfügen
            try:
                tables = page.find_tables()
                for table in tables:
                    extracted_table = table.extract()
                    if extracted_table and len(extracted_table) > 0:
                        num_rows = len(extracted_table)
                        num_cols = len(extracted_table[0])

                        word_table = word_doc.add_table(rows=num_rows, cols=num_cols)
                        for row_idx, row in enumerate(extracted_table):
                            for col_idx, cell in enumerate(row):
                                word_table.cell(row_idx, col_idx).text = str(cell) if cell else ""
            except Exception as table_error:
                print(f"⚠Warnung: Fehler beim Extrahieren einer Tabelle auf Seite {i}: {table_error}")

            # Bilder & Diagramme extrahieren und ins Word-Dokument einfügen
            for img_index, img in enumerate(page.get_images(full=True), start=1):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                
                # Temporäre Datei für das Bild erstellen
                with NamedTemporaryFile(delete=False, suffix=f".{image_ext}") as temp_img:
                    temp_img.write(image_bytes)
                    temp_img_path = temp_img.name
                
                # Bild in Word einfügen
                word_doc.add_paragraph(f"Bild {img_index} von Seite {i}:")
                word_doc.add_picture(temp_img_path, width=Inches(5))
                
                # Temporäre Datei löschen
                os.remove(temp_img_path)

            # Falls die Seite Diagramme enthält, explizit darauf hinweisen
            word_doc.add_paragraph(f"Falls Diagramme auf Seite {i} nicht übernommen wurden, bitte prüfen.")

            # Seitenumbruch nach jeder PDF-Seite
            if i < len(doc):
                word_doc.add_page_break()

        word_doc.save(output_docx)
        print(f"Word-Dokument gespeichert: {output_docx}")

    except Exception as e:
        print(f"Fehler bei der Verarbeitung von {input_pdf}: {e}")

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Nutzung: python script.py <pfad_zur_pdf>")
    else:
        pdf_to_word(sys.argv[1])
